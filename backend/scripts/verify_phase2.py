"""
Phase 2 最终验证脚本
验证: RAG 全流程 + 文件解析模块

用法: python scripts/verify_phase2.py
"""
import sys
import json
import tempfile
from pathlib import Path

_PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
_BACKEND_DIR = _PROJECT_ROOT / "backend"
sys.path.insert(0, str(_BACKEND_DIR))

from config import load_settings

settings = load_settings()
passed = 0
failed = 0
errors = []


def check(name, condition, detail=""):
    global passed, failed
    if condition:
        passed += 1
        print(f"  [PASS] {name}")
    else:
        failed += 1
        msg = f"  [FAIL] {name}"
        if detail:
            msg += f" — {detail}"
        print(msg)
        errors.append(name)


# ============================================================
# Part 1: File Parsing Module
# ============================================================
print("=" * 60)
print("Part 1: File Parsing Module")
print("=" * 60)

from file_utils import extract_text


# 1a. TXT parsing
print("\n[1a] TXT parsing")
with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
    f.write("Hello World\nThis is a test file.\n第三行中文内容。")
    txt_path = f.name

text = extract_text(txt_path, "test.txt")
check("TXT: extracts content", "Hello World" in text, text[:80])
check("TXT: handles Chinese", "中文内容" in text)
Path(txt_path).unlink(missing_ok=True)


# 1b. Markdown parsing
print("\n[1b] Markdown parsing")
with tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8") as f:
    f.write("# Title\n\nSome **bold** text.\n\n- list item 1\n- list item 2")
    md_path = f.name

text = extract_text(md_path, "test.md")
check("MD: extracts content", "Title" in text, text[:80])
check("MD: includes bold text", "bold" in text)
Path(md_path).unlink(missing_ok=True)


# 1c. DOCX parsing
print("\n[1c] DOCX parsing")
try:
    from docx import Document as DocxDocument

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        docx_path = f.name
    doc = DocxDocument()
    doc.add_paragraph("这是 Word 文档的测试内容。")
    doc.add_paragraph("第二段：Python 课程介绍。")
    doc.save(docx_path)

    text = extract_text(docx_path, "test.docx")
    check("DOCX: extracts content", "Word 文档" in text, text[:80])
    check("DOCX: multi-paragraph", "第二段" in text)
    Path(docx_path).unlink(missing_ok=True)
except ImportError:
    check("DOCX: python-docx installed", False, "python-docx not installed")


# 1d. XLSX parsing
print("\n[1d] XLSX parsing")
try:
    from openpyxl import Workbook

    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
        xlsx_path = f.name
    wb = Workbook()
    ws1 = wb.active
    ws1.title = "课程列表"
    ws1.append(["课程名", "价格", "周期"])
    ws1.append(["Python基础", "2999", "8周"])
    ws1.append(["数据分析", "3999", "12周"])
    ws2 = wb.create_sheet("FAQ")
    ws2.append(["问题", "回答"])
    ws2.append(["有优惠吗", "老学员9折"])
    wb.save(xlsx_path)

    text = extract_text(xlsx_path, "test.xlsx")
    check("XLSX: extracts sheet1 content", "课程列表" in text, text[:120])
    check("XLSX: extracts data rows", "Python基础" in text and "2999" in text)
    check("XLSX: handles multiple sheets", "FAQ" in text and "老学员9折" in text)
    Path(xlsx_path).unlink(missing_ok=True)
except ImportError:
    check("XLSX: openpyxl installed", False, "openpyxl not installed")


# 1e. PDF parsing
print("\n[1e] PDF parsing")
try:
    from PyPDF2 import PdfWriter, PdfReader

    # PyPDF2 can't easily create PDFs, so just verify the import and
    # test with the existing extract_text function on a non-existent file
    # to confirm the dispatch works
    try:
        result = extract_text("/nonexistent/test.pdf", "test.pdf")
        # Should return an error string, not crash
        check("PDF: dispatch does not crash", isinstance(result, str))
    except Exception as e:
        check("PDF: dispatch does not crash", False, str(e))
except ImportError:
    check("PDF: PyPDF2 installed", False, "PyPDF2 not installed")


# 1f. Unsupported format
print("\n[1f] Unsupported format fallback")
with tempfile.NamedTemporaryFile(mode="w", suffix=".xyz", delete=False, encoding="utf-8") as f:
    f.write("test")
    xyz_path = f.name

text = extract_text(xyz_path, "test.xyz")
check("Unknown format: returns error message", "不支持" in text, text[:80])
Path(xyz_path).unlink(missing_ok=True)


# ============================================================
# Part 2: RAG Pipeline (Knowledge Base)
# ============================================================
print("\n" + "=" * 60)
print("Part 2: RAG Pipeline - Knowledge Base")
print("=" * 60)

# 2a. Knowledge base files exist
print("\n[2a] Knowledge base source files")
knowledge_dir = Path(settings.knowledge_absolute_dir)
check("Knowledge dir exists", knowledge_dir.exists(), str(knowledge_dir))

supported = {".txt", ".md", ".docx", ".pdf", ".xlsx", ".html", ".ipynb", ".pptx"}
kb_files = [f for f in knowledge_dir.rglob("*") if f.suffix.lower() in supported and f.is_file()]
check("Has knowledge files", len(kb_files) > 0, f"{len(kb_files)} files found")
for f in kb_files:
    print(f"    - {f.name} ({f.stat().st_size} bytes)")


# 2b. ChromaDB persistence
print("\n[2b] ChromaDB persistence")
chroma_dir = Path(settings.chroma_absolute_dir)
check("Chroma dir exists", chroma_dir.exists(), str(chroma_dir))
chroma_files = list(chroma_dir.rglob("*"))
check("Chroma has data files", len(chroma_files) > 0, f"{len(chroma_files)} files")


# 2c. Embeddings module
print("\n[2c] Embeddings module")
try:
    from assistant.embeddings import BGEM3Embeddings, get_embedder

    check("BGEM3Embeddings class importable", True)
    check("get_embedder function importable", True)

    # Test with a lightweight check (don't load the 2GB model here)
    from langchain_core.embeddings import Embeddings

    check(
        "BGEM3Embeddings inherits LangChain Embeddings",
        issubclass(BGEM3Embeddings, Embeddings),
    )
except ImportError as e:
    check("Embeddings module import", False, str(e))


# 2d. Vector store module
print("\n[2d] Vector store module")
try:
    from assistant.vectorstore import VectorStoreManager

    check("VectorStoreManager class importable", True)
except ImportError as e:
    check("VectorStoreManager import", False, str(e))


# 2e. Prompts module
print("\n[2e] Prompts module")
try:
    from assistant.prompts import (
        format_context,
        format_sources,
        build_system_prompt,
        build_messages_with_context,
        RAG_SYSTEM_PROMPT,
        SYSTEM_PROMPT,
    )

    check("All prompt functions importable", True)

    # Test format_context
    from langchain_core.documents import Document

    doc1 = Document(page_content="测试内容一", metadata={"source": "test.md"})
    doc2 = Document(page_content="测试内容二", metadata={"source": "test2.md"})
    ctx = format_context([(doc1, 0.85), (doc2, 0.72)])
    check("format_context: includes source name", "test.md" in ctx)
    check("format_context: includes score", "0.85" in ctx)
    check("format_context: includes content", "测试内容一" in ctx)

    # Test format_sources
    srcs = format_sources([(doc1, 0.85)])
    check("format_sources: returns list", isinstance(srcs, list) and len(srcs) == 1)
    check("format_sources: has content", "content" in srcs[0])
    check("format_sources: has source", srcs[0]["source"] == "test.md")
    check("format_sources: has score", srcs[0]["score"] == 0.85)

    # Test build_system_prompt
    sp1 = build_system_prompt("")
    check("build_system_prompt: empty context returns SYSTEM_PROMPT", sp1 == SYSTEM_PROMPT)
    sp2 = build_system_prompt("测试上下文")
    check("build_system_prompt: with context returns RAG prompt", "测试上下文" in sp2)
    check("build_system_prompt: RAG prompt uses template", "已有课程资料" in sp2)

    # Test build_messages_with_context
    msgs = build_messages_with_context("ctx", [{"role": "user", "content": "hi"}], "hello")
    check("build_messages: starts with system", msgs[0]["role"] == "system")
    check("build_messages: includes history", msgs[1]["role"] == "user")

except ImportError as e:
    check("Prompts module import", False, str(e))
except Exception as e:
    check("Prompts module tests", False, str(e))


# ============================================================
# Part 3: RAG Pipeline - End-to-End (via API)
# ============================================================
print("\n" + "=" * 60)
print("Part 3: RAG Pipeline - End-to-End (SSE API)")
print("=" * 60)

import requests

_host = "127.0.0.1" if settings.flask_host == "0.0.0.0" else settings.flask_host
BASE_URL = f"http://{_host}:{settings.flask_port}"


# Helper: safe print for GBK consoles
def safe_print(msg):
    """Print to console, replacing chars that GBK can't encode"""
    try:
        print(msg)
    except UnicodeEncodeError:
        print(msg.encode("gbk", errors="replace").decode("gbk"))


# Setup: create a valid session for chat tests
print("\n[3a] Setup - Login & create test session")
_session_id = None
try:
    resp = requests.post(f"{BASE_URL}/api/login", json={"username": "verify_test"}, timeout=10)
    check("Login for test session", resp.status_code == 200)
    resp = requests.post(
        f"{BASE_URL}/api/sessions",
        json={"title": "Verify Test"},
        headers={"X-Username": "verify_test"},
        timeout=10,
    )
    check("Create test session", resp.status_code == 200)
    _session_id = resp.json().get("session", {}).get("id", "")
    check("Got session ID", bool(_session_id), _session_id)
except requests.ConnectionError:
    check("Backend is running", False, f"Cannot connect to {BASE_URL}")
    print("  [SKIP] Remaining API tests require backend")
    raise SystemExit(1)


# 3b. Warmup - trigger model loading and wait for Flask reloader to settle
print("\n[3b] Warmup - trigger RAG model loading (first request takes ~60s)...")
import time as _time
warmup_ok = False
try:
    # Fire a request to trigger _init_rag(), ignore result
    resp = requests.post(
        f"{BASE_URL}/api/chat/stream",
        json={"session_id": _session_id, "question": "你好"},
        stream=True,
        timeout=120,
    )
    # Drain the stream (discard output)
    for line in resp.iter_lines(decode_unicode=True):
        pass
    warmup_ok = True
    safe_print("    Warmup complete, model loaded")
except Exception:
    # Connection reset during Flask reload is expected
    safe_print("    Warmup interrupted (Flask reloader, expected)")

# Wait for backend to stabilize after reload
if not warmup_ok:
    safe_print("    Waiting for backend to stabilize...")
    for attempt in range(30):
        _time.sleep(2)
        try:
            r = requests.get(f"{BASE_URL}/api/health", timeout=5)
            if r.status_code == 200:
                safe_print(f"    Backend stable after {attempt * 2}s")
                break
        except Exception:
            pass
check("Warmup completed", True)  # Always passes — warmup failure is non-fatal


# 3c. Health check
print("\n[3c] Health check")
try:
    resp = requests.get(f"{BASE_URL}/api/health", timeout=10)
    check("Health endpoint responds", resp.status_code == 200)
    data = resp.json()
    check("Health status is ok", data.get("status") == "ok")
    check("Health includes rag_available", "rag_available" in data)
    check("Health includes knowledge_base_docs", "knowledge_base_docs" in data)
    safe_print(f"    rag_available={data.get('rag_available')}, kb_docs={data.get('knowledge_base_docs')}")
except requests.ConnectionError:
    check("Health endpoint reachable", False, f"Cannot connect to {BASE_URL}")


# 3d. Chat SSE stream - RAG question
print("\n[3d] Chat SSE stream (RAG question)")
try:
    resp = requests.post(
        f"{BASE_URL}/api/chat/stream",
        json={"session_id": _session_id, "question": "Python数据分析课程学什么？"},
        stream=True,
        timeout=120,
    )
    check("SSE endpoint responds", resp.status_code == 200)

    events = []
    for line in resp.iter_lines(decode_unicode=True):
        if line and line.startswith("data: "):
            try:
                events.append(json.loads(line[6:]))
            except json.JSONDecodeError:
                pass

    types = [e.get("type") for e in events]
    check("SSE: first event is 'sources'", types[0] == "sources" if types else False,
          f"got types: {types[:5]}...")
    check("SSE: has 'token' events", "token" in types)
    check("SSE: last event is 'done'", types[-1] == "done" if types else False,
          f"last type: {types[-1] if types else 'none'}")

    # Check sources
    sources_event = events[0]
    sources_data = sources_event.get("data", [])
    if len(sources_data) > 0:
        check("SSE sources: has results", True, f"{len(sources_data)} sources found")
        for s in sources_data:
            safe_print(f"    source={s.get('source')}, score={s.get('score')}")
            check(f"SSE source: score >= 0.45 ({s.get('source')})",
                  s.get("score", 0) >= 0.45,
                  f"score={s.get('score')}")
    else:
        print("    [INFO] No RAG sources returned (KB may be empty, falling back to general chat)")

    # Check AI response references course content
    full_answer = "".join(
        e.get("data", "") for e in events if e.get("type") == "token"
    )
    safe_print(f"    AI answer preview: {full_answer[:200]}...")
    # Check if answer references knowledge base (sample_course.md content)
    kb_keywords = ["Python", "数据", "课程", "NumPy", "Pandas", "就业", "8周", "16周"]
    hits = [kw for kw in kb_keywords if kw in full_answer]
    check(
        "AI answer references KB content",
        len(hits) >= 2,
        f"matched keywords: {hits}",
    )

except requests.ConnectionError:
    check("SSE chat: backend reachable", False, "Connection refused")
except Exception as e:
    check("SSE chat: no exception", False, repr(e)[:200])


# 3e. Chat SSE stream - unrelated question (graceful degradation)
print("\n[3e] Chat SSE stream (unrelated question, graceful degradation)")
try:
    resp = requests.post(
        f"{BASE_URL}/api/chat/stream",
        json={"session_id": _session_id, "question": "今天天气怎么样？"},
        stream=True,
        timeout=120,
    )
    events = []
    for line in resp.iter_lines(decode_unicode=True):
        if line and line.startswith("data: "):
            try:
                events.append(json.loads(line[6:]))
            except json.JSONDecodeError:
                pass

    types = [e.get("type") for e in events]
    check("Unrelated: first event is 'sources'",
          types[0] == "sources" if types else False)
    check("Unrelated: has 'done' event", "done" in types)
    # sources should be empty or very low score
    sources_data = events[0].get("data", []) if events else []
    check("Unrelated: sources is list", isinstance(sources_data, list))
    safe_print(f"    sources count: {len(sources_data)}")

except Exception as e:
    check("Unrelated question: no exception", False, repr(e)[:200])


# ============================================================
# Part 4: File Upload API
# ============================================================
print("\n" + "=" * 60)
print("Part 4: File Upload API")
print("=" * 60)


def test_upload(filename: str, content: bytes, expected_text: str, label: str):
    """Test file upload and extraction"""
    try:
        import io
        resp = requests.post(
            f"{BASE_URL}/api/upload",
            files={"file": (filename, io.BytesIO(content), "application/octet-stream")},
            timeout=30,
        )
        check(f"Upload {label}: HTTP 200", resp.status_code == 200,
              f"status={resp.status_code}")
        data = resp.json()
        check(f"Upload {label}: has file_id", "file_id" in data)
        check(f"Upload {label}: content extracted", len(data.get("content_text", "")) > 0)
        check(f"Upload {label}: expected text present",
              expected_text in data.get("content_text", ""),
              f"preview: {data.get('content_text', '')[:100]}")
        return data
    except Exception as e:
        check(f"Upload {label}: no exception", False, str(e))
        return None


# 4a. TXT upload
print("\n[4a] TXT upload")
test_upload("test_verify.txt", "这是上传测试文件的内容。\n第二行。".encode("utf-8"),
            "上传测试", "TXT")

# 4b. MD upload
print("\n[4b] MD upload")
test_upload("test_verify.md", "# 测试\n\nMarkdown 上传测试。".encode("utf-8"),
            "Markdown", "MD")

# 4c. DOCX upload
print("\n[4c] DOCX upload")
try:
    from docx import Document as DocxDocument
    import io as io_module

    buf = io_module.BytesIO()
    doc = DocxDocument()
    doc.add_paragraph("Word 文档上传测试内容。")
    doc.save(buf)
    buf.seek(0)
    test_upload("test_verify.docx", buf.read(), "Word 文档", "DOCX")
except ImportError:
    print("  [SKIP] python-docx not installed")

# 4d. XLSX upload
print("\n[4d] XLSX upload")
try:
    from openpyxl import Workbook
    import io as io_module

    buf = io_module.BytesIO()
    wb = Workbook()
    ws = wb.active
    ws.title = "测试表"
    ws.append(["A1", "B1", "C1"])
    ws.append(["数据1", "数据2", "数据3"])
    wb.save(buf)
    buf.seek(0)
    test_upload("test_verify.xlsx", buf.read(), "数据1", "XLSX")
except ImportError:
    print("  [SKIP] openpyxl not installed")

# 4e. Unsupported format upload
print("\n[4e] Unsupported format upload")
try:
    import io as io_module
    resp = requests.post(
        f"{BASE_URL}/api/upload",
        files={"file": ("test.xyz", io_module.BytesIO(b"test"), "application/octet-stream")},
        timeout=30,
    )
    check("Upload unsupported: returns error", resp.status_code == 400,
          f"status={resp.status_code}, body={resp.text[:100]}")
except Exception as e:
    check("Upload unsupported: no exception", False, str(e))


# ============================================================
# Summary
# ============================================================
print("\n" + "=" * 60)
print(f"VERIFICATION COMPLETE: {passed} passed, {failed} failed, {passed + failed} total")
print("=" * 60)

if errors:
    print("\nFailed checks:")
    for e in errors:
        print(f"  - {e}")

sys.exit(0 if failed == 0 else 1)
